from docx import Document
from docx.shared import Pt

# Initialize document
doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# Contact Info
add_paragraph('Krishna Kumar Raut', bold=True)
doc.add_paragraph('📍 Nepal | 📞 +91-8969135901 | ✉️ rautkrishna2217@gmail.com')
doc.add_paragraph('LinkedIn | GitHub | Portfolio | LeetCode')

# Summary
add_heading('Summary', level=2)
doc.add_paragraph(
    "Detail-oriented Java Backend Developer with 1 year of experience building scalable microservices and RESTful APIs using Spring Boot. "
    "Skilled in optimizing backend performance, writing robust unit/integration tests, and deploying production-grade services. "
    "Proficient with Docker, JWT authentication, and CI/CD workflows. Passionate about backend architecture and delivering high-impact solutions, "
    "seeking to contribute to engineering teams at Amazon or Microsoft."
)

# Technical Skills
add_heading('Technical Skills', level=2)
add_paragraph("Languages: Java, C++, Python, JavaScript, MySQL")
add_paragraph("Frameworks & Tools: Spring Boot, Spring MVC, Spring Data JPA, JUnit, Mockito")
add_paragraph("DevOps & Tools: Docker, Git, GitHub, Redis, CI/CD")
add_paragraph("Concepts: RESTful API Design, Microservices, OOP, SQL Optimization, Authentication & Authorization")

# Experience
add_heading('Experience', level=2)
add_paragraph("Epam Systems – Java Backend Developer", bold=True)
doc.add_paragraph("July 2024 – Present")
doc.add_paragraph("Tech Stack: Java, Spring Boot, Spring Security, Redis, MySQL, JUnit, Docker")
add_bullet("Developed and maintained 10+ RESTful APIs, enabling seamless communication between frontend and backend services.")
add_bullet("Improved backend performance, reducing API response times by 30% via query optimization and Redis caching.")
add_bullet("Wrote 100+ unit and integration tests, achieving 90% code coverage with JUnit and Mockito.")
add_bullet("Implemented secure JWT-based authentication supporting 400+ users with role-based access.")
add_bullet("Utilized Docker and CI/CD pipelines for automated builds and deployments.")

# Projects
add_heading('Projects', level=2)
add_paragraph("Recipe Management System", bold=True)
doc.add_paragraph("Spring Boot, Spring Security, API Gateway, Microservices")
add_bullet("Built a microservices-based application for 500+ users managing 1,000+ recipes.")
add_bullet("Processed 10,000+ daily requests via API Gateway with centralized authentication and load balancing.")
add_bullet("Implemented Spring Security for role-based access; ensured modularity and scalability.")

add_paragraph("Municipality Management System", bold=True)
doc.add_paragraph("HTML, CSS, JS, PHP, MySQL")
add_bullet("Created a record management system managing 5,000+ municipal entries.")
add_bullet("Designed UI and integrated backend using PHP and MySQL for efficient data operations.")
add_bullet("Praised by local authorities for usability and civic impact.")

# Education
add_heading('Education', level=2)
add_paragraph("Vellore Institute of Technology, India", bold=True)
doc.add_paragraph("B.Tech in Computer Science and Engineering (Dec 2020 – July 2024)")
doc.add_paragraph("CGPA: 8.42 / 10")
add_paragraph("National Infotech Secondary School, Nepal", bold=True)
doc.add_paragraph("Higher Secondary - PCM (Apr 2018 – July 2020)")
doc.add_paragraph("CGPA: 3.64 / 4")

# Certifications & Awards
add_heading('Certifications & Awards', level=2)
add_bullet("Full Scholarship Awardee, Indian Embassy – B.Tech at VIT")
add_bullet("Microsoft Engage Program – Selected Mentee")
add_bullet("2★ on CodeChef, 5★ (C++) on HackerRank, 500+ LeetCode problems solved")

# Save document
doc.save("Krishna_Kumar_Raut_Backend_Resume.docx")
